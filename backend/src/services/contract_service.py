from datetime import date
from typing import Optional
from uuid import UUID
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from src.models.all_models import LeaseContract, Property, Tenant, User

async def generate_contract_document(db: AsyncSession, contract_id: UUID) -> Document:
    """Generate a Word document for a lease contract"""
    
    # Fetch contract with relations
    stmt = select(LeaseContract).options(
        selectinload(LeaseContract.property),
        selectinload(LeaseContract.tenant)
    ).where(LeaseContract.id == contract_id)
    
    contract = (await db.execute(stmt)).scalar_one_or_none()
    if not contract:
        raise ValueError("Contract not found")
    
    # Get property owner
    prop_owner = await db.get(User, contract.property.owner_id)
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ДОГОВОР АРЕНДЫ НЕДВИЖИМОГО ИМУЩЕСТВА')
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    
    # Contract number and date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f'№ {contract.number}\n')
    date_para.add_run(f'от «{contract.start_date.day}» {_get_month_name(contract.start_date.month)} {contract.start_date.year} г.')
    
    doc.add_paragraph()  # Empty line
    
    # Introduction
    intro = doc.add_paragraph()
    intro.add_run('Настоящим договор заключен между:').bold = True
    
    # Landlord section
    landlord_para = doc.add_paragraph(style='List Bullet')
    landlord_para.add_run('Арендодатель: ').bold = True
    landlord_para.add_run(
        f'{prop_owner.name or prop_owner.username} '
        f'(телефон: {prop_owner.phone or "не указан"}), '
        f'далее именуемый "Арендодатель"'
    )
    
    # Tenant section
    tenant_para = doc.add_paragraph(style='List Bullet')
    tenant_para.add_run('Арендатор: ').bold = True
    tenant_para.add_run(
        f'{contract.tenant.name} '
        f'(телефон: {contract.tenant.phone}), '
        f'далее именуемый "Арендатор"'
    )
    
    doc.add_paragraph()  # Empty line
    
    # 1. SUBJECT OF CONTRACT
    doc.add_paragraph('1. ПРЕДМЕТ ДОГОВОРА', style='Heading 2')
    
    subject = doc.add_paragraph()
    subject.add_run(
        f'1.1. Арендодатель передает, а Арендатор принимает во временное владение '
        f'и пользование недвижимое имущество (далее - "Объект"), '
        f'расположенное по адресу:\n'
    )
    
    address_para = doc.add_paragraph(style='List Bullet')
    address_para.add_run(
        f'Город: {contract.property.city}\n'
        f'Улица: {contract.property.street}, дом {contract.property.house}'
    )
    if contract.property.unit:
        address_para.add_run(f', кв. {contract.property.unit}')
    
    property_details = doc.add_paragraph()
    property_details.add_run('\n1.2. Характеристики объекта:\n').bold = True
    
    details_table = doc.add_table(rows=6, cols=2)
    details_table.style = 'Light Grid Accent 1'
    
    details = [
        ('Тип объекта:', _get_property_type_name(contract.property.type)),
        ('Площадь:', f'{contract.property.area} м²'),
        ('Количество комнат:', str(contract.property.rooms or '-')),
        ('Этаж:', f'{contract.property.floor}/{contract.property.total_floors}' if contract.property.floor else '-'),
        ('Описание:', contract.property.description or 'Нет'),
    ]
    
    for idx, (key, value) in enumerate(details, 1):
        details_table.rows[idx].cells[0].text = key
        details_table.rows[idx].cells[1].text = value
    
    doc.add_paragraph()  # Empty line
    
    # 2. TERM OF CONTRACT
    doc.add_paragraph('2. СРОКИ ДЕЙСТВИЯ ДОГОВОРА', style='Heading 2')
    
    term_para = doc.add_paragraph()
    term_para.add_run(
        f'2.1. Договор заключается на период с «{contract.start_date.day}» '
        f'{_get_month_name(contract.start_date.month)} {contract.start_date.year} года '
        f'по «{contract.end_date.day}» {_get_month_name(contract.end_date.month)} {contract.end_date.year} года.\n'
    )
    
    duration = (contract.end_date - contract.start_date).days // 30
    term_para.add_run(f'2.2. Общая продолжительность аренды: {duration} месяцев.\n')
    term_para.add_run(f'2.3. Периодичность платежей: {_get_payment_period_name(contract.payment_period)}.')
    
    doc.add_paragraph()  # Empty line
    
    # 3. RENT AND PAYMENTS
    doc.add_paragraph('3. РАЗМЕР ПЛАТЕЖА И ПОРЯДОК РАСЧЕТОВ', style='Heading 2')
    
    payment_para = doc.add_paragraph()
    payment_para.add_run(f'3.1. Размер арендной платы: {contract.rent_amount} BYN\n')
    payment_para.add_run(f'3.2. Периодичность платежа: {_get_payment_period_name(contract.payment_period)}\n')
    payment_para.add_run(f'3.3. Платеж должен быть произведен не позднее 5-го числа расчетного периода.\n')
    
    if contract.deposit_amount > 0:
        payment_para.add_run(
            f'3.4. Сумма залога (депозита): {contract.deposit_amount} BYN\n'
        )
        payment_para.add_run(
            f'3.5. Залог подлежит возврату в течение 30 дней после окончания договора, '
            f'при условии отсутствия повреждений и задолженности.\n'
        )
    
    doc.add_paragraph()  # Empty line
    
    # 4. OBLIGATIONS
    doc.add_paragraph('4. ПРАВА И ОБЯЗАННОСТИ СТОРОН', style='Heading 2')
    
    obligations = doc.add_paragraph()
    obligations.add_run('4.1. Обязанности Арендатора:\n').bold = True
    
    tenant_obligations = [
        'Своевременно вносить арендную плату',
        'Содержать объект в надлежащем состоянии',
        'Не сдавать объект в субаренду без согласия Арендодателя',
        'Не производить перепланировку без письменного согласия',
        'Соблюдать правила внутреннего распорядка',
    ]
    
    for obligation in tenant_obligations:
        doc.add_paragraph(obligation, style='List Bullet')
    
    doc.add_paragraph()
    landlord_obligations = doc.add_paragraph()
    landlord_obligations.add_run('4.2. Обязанности Арендодателя:\n').bold = True
    
    owner_obligations = [
        'Передать объект в надлежащем состоянии',
        'Обеспечить возможность беспрепятственного пользования объектом',
        'Проводить текущий ремонт общего имущества',
        'Гарантировать мирное владение',
    ]
    
    for obligation in owner_obligations:
        doc.add_paragraph(obligation, style='List Bullet')
    
    doc.add_paragraph()  # Empty line
    
    # 5. TERMINATION
    doc.add_paragraph('5. РАСТОРЖЕНИЕ ДОГОВОРА', style='Heading 2')
    
    termination = doc.add_paragraph()
    termination.add_run(
        '5.1. Договор может быть расторгнут досрочно в случаях:\n'
    )
    
    termination_reasons = [
        'По соглашению сторон',
        'По письменному уведомлению за 30 дней',
        'В случае неуплаты в течение 15 дней после срока платежа',
        'При существенном нарушении условий договора',
    ]
    
    for reason in termination_reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_paragraph()  # Empty line
    
    # 6. SIGNATURES
    doc.add_paragraph('6. ПОДПИСИ СТОРОН', style='Heading 2')
    
    signatures = doc.add_table(rows=3, cols=2)
    signatures.style = 'Table Grid'
    
    signatures.rows[0].cells[0].text = 'АРЕНДОДАТЕЛЬ'
    signatures.rows[0].cells[1].text = 'АРЕНДАТОР'
    
    signatures.rows[1].cells[0].text = '_________________________'
    signatures.rows[1].cells[1].text = '_________________________'
    
    signatures.rows[2].cells[0].text = prop_owner.name or prop_owner.username
    signatures.rows[2].cells[1].text = contract.tenant.name
    
    return doc


def _get_month_name(month: int) -> str:
    """Get Russian month name"""
    months = {
        1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
        5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
        9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
    }
    return months.get(month, '')


def _get_property_type_name(prop_type: str) -> str:
    """Get Russian property type name"""
    types = {
        'flat': 'Квартира',
        'office': 'Офис',
        'warehouse': 'Склад',
        'other': 'Прочее'
    }
    return types.get(prop_type, prop_type)


def _get_payment_period_name(period: str) -> str:
    """Get Russian payment period name"""
    periods = {
        'month': 'Ежемесячно',
        'quarter': 'Ежеквартально'
    }
    return periods.get(period, period)
